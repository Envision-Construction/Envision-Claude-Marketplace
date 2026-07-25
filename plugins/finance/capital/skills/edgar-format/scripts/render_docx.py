#!/usr/bin/env python3
"""
render_docx.py — render extracted JSON to an EDGAR-styled .docx.

DOCX is an edit-first format — used when the recipient marks up the document
(redlines, comments) rather than reading a static page. The forensic spec
applies the same way as in xlsx/html:

  - Body font: Times New Roman 10pt
  - Title block: registrant (11pt bold centered), title (10pt bold ALL CAPS
    centered), units note (9pt italic centered)
  - Statement table: right-aligned numerics, em-dash for zero, parens for
    negatives, single-rule top border on subtotals, single-rule top +
    double-rule bottom on grand totals
  - Floating $ on first data row + every subtotal/grand_total row
  - EPS rows: 2-decimal; share counts: integer
  - No body bold — emphasis is borders, not weight

Usage:
    python3 render_docx.py <extracted.json> <output.docx>
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover
    raise SystemExit(f"python-docx required for docx rendering: {exc}")


FONT = "Times New Roman"


def _is_eps_label(label: str) -> bool:
    l = label.lower()
    return "per share" in l or l.strip() in {"basic", "diluted"}


def _is_share_count_label(label: str) -> bool:
    l = label.lower()
    return "weighted-average" in l or "weighted average" in l or "shares outstanding" in l


def _fmt_int(v: float) -> str:
    if v < 0:
        return f"({int(round(-v)):,})"
    return f"{int(round(v)):,}"


def _fmt_decimal(v: float) -> str:
    if v < 0:
        return f"({-v:,.2f})"
    return f"{v:,.2f}"


def _set_cell_border(cell, *, top: str = "", bottom: str = "") -> None:
    """Apply top/bottom border to a docx table cell via raw OOXML.
    `top`/`bottom` can be "single" (thin) or "double". Empty string skips."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, style in [("top", top), ("bottom", bottom)]:
        if not style:
            continue
        node = tcBorders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tcBorders.append(node)
        node.set(qn("w:val"), style)
        node.set(qn("w:sz"), "8" if style == "double" else "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def _run(para, text: str, *, size_pt: int = 10, bold: bool = False, italic: bool = False) -> None:
    r = para.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size_pt)
    r.font.bold = bold
    r.font.italic = italic


def _render_segment(doc: Document, segment: dict[str, Any], metadata: dict[str, Any]) -> None:
    period_labels: list[str] = segment.get("period_labels", []) or []
    n_periods = max(len(period_labels), 1)
    rows = segment.get("rows", []) or []

    # ----- Title block ---------------------------------------------------
    registrant = (metadata.get("registrant") or "").strip()
    if registrant:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, registrant, size_pt=11, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, segment.get("title", "").upper(), size_pt=10, bold=True)

    units = segment.get("units_note", "")
    if units:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, units, size_pt=9, italic=True)
        p.paragraph_format.space_after = Pt(12)

    # ----- Period header -------------------------------------------------
    period_header = segment.get("period_header", "")
    if period_header:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, period_header, size_pt=10)

    # ----- Table ---------------------------------------------------------
    table = doc.add_table(rows=1 + len(rows), cols=1 + n_periods)
    table.autofit = False

    # Column widths — label wide, numerics narrower
    table.columns[0].width = Inches(4.0)
    for i in range(1, 1 + n_periods):
        table.columns[i].width = Inches(1.1)

    # Header row (column years)
    header_row = table.rows[0]
    header_row.cells[0].text = ""
    for i, label in enumerate(period_labels or [""]):
        cell = header_row.cells[1 + i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(p, label, size_pt=10, bold=True)
        _set_cell_border(cell, bottom="single")

    # Body rows
    first_data_row_written = False
    current_section = ""
    for row_idx, r in enumerate(rows, start=1):
        label = r.get("label", "") or ""
        role = r.get("role", "data")
        indent = int(r.get("indent", 0) or 0)
        values = r.get("values", []) or []
        tr = table.rows[row_idx]

        # Label cell
        label_cell = tr.cells[0]
        label_cell.text = ""
        lp = label_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if indent > 0:
            lp.paragraph_format.left_indent = Inches(0.25 * min(indent, 3))
        display = label.upper() if role == "banner" else label
        _run(lp, display, size_pt=10)

        if role == "spacer":
            continue
        if role == "section_header":
            current_section = label.lower()
            continue
        if role in {"banner", "text_only"}:
            continue

        # Decide formatter
        section_is_shares = "shares outstanding" in current_section or "weighted-average" in current_section
        section_is_eps = "per share" in current_section or "earnings per" in current_section
        if section_is_shares or _is_share_count_label(label):
            fmt = _fmt_int
        elif section_is_eps or _is_eps_label(label):
            fmt = _fmt_decimal
        else:
            fmt = _fmt_int

        floats_dollar = role in {"subtotal", "grand_total"} or not first_data_row_written

        for i, v in enumerate(values):
            cell = tr.cells[1 + i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if v is None:
                _run(p, "—", size_pt=10)
                content = "—"
            else:
                prefix = "$ " if floats_dollar else ""
                _run(p, prefix + fmt(float(v)), size_pt=10)

            if role == "subtotal":
                _set_cell_border(cell, top="single")
            elif role == "grand_total":
                _set_cell_border(cell, top="single", bottom="double")

        if role == "data" and any(v is not None for v in values):
            first_data_row_written = True

    # Extra blank paragraph after the table for breathing room
    doc.add_paragraph()


def render(extracted: dict[str, Any], out_path: Path) -> None:
    doc = Document()
    # Set page margins per spec — Letter, 1" top/bottom, 0.75" left/right
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Set the default style to Times New Roman 10pt
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)

    metadata = extracted.get("metadata") or {}
    segments = extracted.get("segments") or []
    if not segments:
        p = doc.add_paragraph()
        _run(p,
             "edgar-format render_docx.py: no statement segments detected. "
             "For narrative full-filing output, the docx renderer currently "
             "only emits an empty document — use render_html.py or render_pdf.py "
             "for a cover-page-and-body filing.",
             size_pt=10)
    else:
        for seg in segments:
            _render_segment(doc, seg, metadata)

    doc.save(out_path)


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("extracted", type=Path)
    ap.add_argument("output", type=Path)
    args = ap.parse_args(argv)
    data = json.loads(args.extracted.read_text(encoding="utf-8"))
    render(data, args.output)
    print(f"wrote {args.output}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
