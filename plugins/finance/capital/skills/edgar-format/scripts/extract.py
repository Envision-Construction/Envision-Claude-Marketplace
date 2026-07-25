#!/usr/bin/env python3
"""
extract.py — universal input → structured EDGAR-renderable JSON.

Reads any supported file format (xlsx, xlsm, csv, tsv, docx, md, txt, html, pdf,
json) and produces a single JSON document on stdout (or to --out FILE) that the
renderer scripts in this skill consume.

Statement-type detection runs over the extracted content and sets `mode_hint` to
one of: segment:income_statement | segment:balance_sheet | segment:cash_flows |
segment:stockholders_equity | segment:comprehensive_income | full_filing |
ambiguous.

Why this exists: every renderer needs the same shape regardless of input
format. Centralizing extraction + detection here keeps the renderers focused on
applying EDGAR style, not parsing.

Usage:
    python3 extract.py <input_file> [--out <output.json>] [--mode <hint>]
                                   [--registrant NAME] [--filing-type 10-K]
                                   [--period-end YYYY-MM-DD]
"""
from __future__ import annotations

import argparse
import csv
import json
import re
import sys
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Iterable

# ---------------------------------------------------------------------------
# Detection heuristics — keyed to references/financial-statements.md
# ---------------------------------------------------------------------------

_SIGNATURES: dict[str, tuple[str, ...]] = {
    "income_statement": (
        "revenue", "net sales", "cost of revenue", "cost of goods sold",
        "gross profit", "operating expenses", "income from operations",
        "operating income", "net income", "earnings per share", "diluted",
        "weighted-average shares", "weighted average shares",
        "selling, general and administrative", "research and development",
    ),
    "balance_sheet": (
        "total current assets", "total assets", "total liabilities",
        "stockholders' equity", "stockholders equity", "accumulated deficit",
        "retained earnings", "accounts receivable", "accounts payable",
        "property and equipment", "goodwill", "current portion of",
        "additional paid-in capital",
    ),
    "cash_flows": (
        "cash flows from operating activities",
        "cash flows from investing activities",
        "cash flows from financing activities",
        "net cash provided by", "net cash used in", "supplemental disclosures",
        "net increase in cash", "net decrease in cash",
        "depreciation and amortization", "stock-based compensation",
    ),
    "stockholders_equity": (
        "balance, december 31", "balance, beginning of year",
        "balance at beginning", "balance, end of year",
        "additional paid-in capital", "treasury stock",
        "accumulated other comprehensive",
    ),
    "comprehensive_income": (
        "other comprehensive income", "foreign currency translation",
        "comprehensive income", "available-for-sale securities",
        "unrealized gains", "unrealized losses",
    ),
}


def _score_statement_type(text_lower: str) -> dict[str, int]:
    """Count signature label hits per type. Highest score wins."""
    return {
        stype: sum(1 for sig in sigs if sig in text_lower)
        for stype, sigs in _SIGNATURES.items()
    }


def _detect_mode(text: str) -> str:
    """Return mode_hint for the extracted text."""
    text_lower = text.lower()
    scores = _score_statement_type(text_lower)
    top_score = max(scores.values()) if scores else 0
    if top_score < 2:
        return "ambiguous" if top_score >= 1 else "full_filing"
    winners = [s for s, c in scores.items() if c == top_score]
    if len(winners) == 1:
        return f"segment:{winners[0]}"
    # Tie-break: comprehensive_income often co-occurs with income_statement,
    # cash_flows co-occurs with balance_sheet language. Prefer the most
    # statement-specific signature.
    priority = (
        "cash_flows", "stockholders_equity", "comprehensive_income",
        "balance_sheet", "income_statement",
    )
    for p in priority:
        if p in winners:
            return f"segment:{p}"
    return "ambiguous"


# ---------------------------------------------------------------------------
# Data model — what the renderers consume
# ---------------------------------------------------------------------------

@dataclass
class StatementRow:
    label: str
    values: list[float | None] = field(default_factory=list)
    indent: int = 0
    role: str = "data"  # data | section_header | subtotal | grand_total | banner | spacer | text_only


@dataclass
class StatementSegment:
    statement_type: str            # income_statement | balance_sheet | ...
    title: str                     # canonical EDGAR title
    units_note: str                # "(In thousands, except per share data)"
    period_header: str             # "Years Ended December 31," or "December 31,"
    period_labels: list[str]       # ["2026", "2025", "2024"] or dates for BS
    rows: list[StatementRow] = field(default_factory=list)


@dataclass
class FilingMetadata:
    registrant: str | None = None
    filing_type: str | None = None
    period_end: str | None = None
    consolidated: bool = False
    state_of_incorporation: str | None = None
    irs_ein: str | None = None
    commission_file_number: str | None = None
    address: str | None = None
    telephone: str | None = None


@dataclass
class Extracted:
    mode_hint: str
    source_format: str
    metadata: FilingMetadata = field(default_factory=FilingMetadata)
    segments: list[StatementSegment] = field(default_factory=list)
    body_sections: list[dict[str, Any]] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Format-specific extractors
# ---------------------------------------------------------------------------

_TITLE_HINT_MAP = {
    "income_statement": "CONSOLIDATED STATEMENTS OF OPERATIONS",
    "balance_sheet": "CONSOLIDATED BALANCE SHEETS",
    "cash_flows": "CONSOLIDATED STATEMENTS OF CASH FLOWS",
    "stockholders_equity": "CONSOLIDATED STATEMENTS OF STOCKHOLDERS' EQUITY",
    "comprehensive_income": "CONSOLIDATED STATEMENTS OF COMPREHENSIVE INCOME",
}

_DEFAULT_UNITS = {
    "income_statement": "(In thousands, except per share data)",
    "balance_sheet": "(In thousands, except share and per share data)",
    "cash_flows": "(In thousands)",
    "stockholders_equity": "(In thousands, except share data)",
    "comprehensive_income": "(In thousands)",
}

_DEFAULT_PERIOD_HEADER = {
    "income_statement": "Years Ended December 31,",
    "balance_sheet": "December 31,",
    "cash_flows": "Years Ended December 31,",
    "stockholders_equity": "",  # equity statement uses Balance rows
    "comprehensive_income": "Years Ended December 31,",
}


def _parse_number(val: Any) -> float | None:
    """Convert a cell value to float; return None for blanks / em-dashes."""
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return float(val)
    s = str(val).strip()
    if s == "" or s in {"—", "-", "–", "N/A", "NM", "nm"}:
        return None
    # Strip $ , spaces; convert (123) → -123
    neg = s.startswith("(") and s.endswith(")")
    s = s.replace("$", "").replace(",", "").replace(" ", "")
    if neg:
        s = "-" + s.strip("()")
    try:
        return float(s)
    except ValueError:
        return None


def _classify_row(label: str, indent: int, statement_type: str = "") -> str:
    """Classify a row based on its label. `statement_type` disambiguates rows
    whose role depends on context — e.g. "Net income" is a grand_total on the
    income statement but the entry-point data line on the cash flow statement.
    """
    l = label.lower().strip()
    if not l:
        return "spacer"
    if l in {"assets", "liabilities", "liabilities and stockholders' equity",
             "liabilities and stockholders equity"}:
        return "banner"
    if l.endswith(":"):
        return "section_header"
    if l.startswith("commitments and contingencies"):
        return "text_only"

    if statement_type == "income_statement":
        if l in {"net income", "net loss"}:
            return "grand_total"
        if l.startswith(("total ", "gross profit", "income from operations",
                         "operating income", "income before income taxes",
                         "income before taxes")):
            return "subtotal"
        return "data"

    if statement_type == "balance_sheet":
        if l.startswith(("total assets", "total liabilities and stockholders")):
            return "grand_total"
        if l.startswith("total "):
            return "subtotal"
        return "data"

    if statement_type == "cash_flows":
        # The ending cash balance is the only grand total on a CFS — Net
        # income at the top of the operating activities reconciliation is
        # just data here, NOT a grand total (that's an income-statement role).
        if l.startswith("cash and cash equivalents, end") \
           or l.startswith("cash, end of period"):
            return "grand_total"
        if l.startswith(("net cash provided by", "net cash used in",
                         "net cash provided by (used in)",
                         "net increase in cash", "net decrease in cash",
                         "net change in cash")):
            return "subtotal"
        return "data"

    if statement_type == "comprehensive_income":
        if l == "comprehensive income":
            return "grand_total"
        if l.startswith("total other comprehensive"):
            return "subtotal"
        return "data"

    if statement_type == "stockholders_equity":
        # Renderer is responsible for promoting the FINAL Balance row to
        # grand_total — at classify time we don't know which is last.
        if l.startswith(("balance, ", "balance at ", "balance as of ")):
            return "subtotal"
        return "data"

    # Unknown statement_type — fall back to generic heuristics
    if l in {"net income", "net loss", "comprehensive income"} \
       or l.startswith(("total assets", "total liabilities and stockholders",
                        "cash and cash equivalents, end")):
        return "grand_total"
    if l.startswith(("total ", "gross profit", "net cash provided by",
                     "net cash used in", "net cash provided by (used in)",
                     "net increase in cash", "net decrease in cash",
                     "income from operations", "income before income taxes")):
        return "subtotal"
    return "data"


def _from_xlsx(path: Path) -> tuple[str, list[StatementSegment]]:
    """Read an xlsx workbook. Returns (joined_text_for_detection, segments)."""
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise SystemExit(f"openpyxl required for xlsx input: {exc}")

    wb = load_workbook(path, data_only=True)
    all_text_parts: list[str] = []
    segments: list[StatementSegment] = []

    for sheet in wb.worksheets:
        sheet_text: list[str] = []
        raw_rows: list[list[Any]] = []
        for row in sheet.iter_rows(values_only=True):
            cells = [c for c in row]
            if all(c is None or str(c).strip() == "" for c in cells):
                raw_rows.append([])
                continue
            raw_rows.append(list(cells))
            for cell in cells:
                if cell is not None:
                    sheet_text.append(str(cell))
        joined = "\n".join(sheet_text)
        all_text_parts.append(joined)

        # Detect statement type for this sheet
        mode = _detect_mode(joined)
        if not mode.startswith("segment:"):
            continue
        stype = mode.split(":", 1)[1]

        seg = _segment_from_rows(stype, raw_rows)
        if seg.rows:
            segments.append(seg)

    return "\n".join(all_text_parts), segments


def _segment_from_rows(stype: str, raw_rows: list[list[Any]]) -> StatementSegment:
    """Convert a 2D grid into a StatementSegment."""
    # Find the numeric column boundary: assume label is column 0, numbers are
    # the remaining columns. Period labels are the first row that contains
    # year-like or date-like values in those columns.
    period_labels: list[str] = []
    data_start = 0
    for i, row in enumerate(raw_rows):
        if len(row) < 2:
            continue
        right = row[1:]
        numeric_or_year = sum(
            1 for c in right
            if c is not None and (
                _parse_number(c) is not None
                or re.match(r"^\s*(19|20)\d{2}\s*$", str(c))
                or re.match(r"[A-Za-z]+ ?\d{1,2}, ?\d{4}", str(c))
            )
        )
        if numeric_or_year >= max(1, len([c for c in right if c is not None]) - 1):
            # Probable header row OR first data row. A header row is one where
            # EVERY non-null cell matches a year ("2026") or date ("December
            # 31, 2026") token — capture those as period_labels. Otherwise
            # we've already passed the header and this is the first data row.
            non_null = [c for c in right if c is not None]
            def _looks_like_period(c: Any) -> bool:
                s = str(c).strip()
                return bool(
                    re.match(r"^(19|20)\d{2}$", s)
                    or re.match(r"^[A-Za-z]+ ?\d{1,2}, ?\d{4}$", s)
                )
            if non_null and all(_looks_like_period(c) for c in non_null):
                period_labels = [str(c).strip() for c in non_null]
                data_start = i + 1
                break
            data_start = i
            break

    rows_out: list[StatementRow] = []
    for r in raw_rows[data_start:]:
        if not r:
            rows_out.append(StatementRow(label="", role="spacer"))
            continue
        label_raw = r[0]
        if label_raw is None:
            label = ""
        else:
            label = str(label_raw).rstrip()
        # Estimate indent from leading whitespace. EDGAR source files vary —
        # 2 spaces, 3 spaces, half-tab, full tab. Treat any leading whitespace
        # as at least one indent level; deepen on multiples of two.
        stripped = label.lstrip()
        indent_chars = len(label) - len(stripped)
        if indent_chars == 0:
            indent = 0
        else:
            indent = min(3, 1 + (indent_chars // 2 - 1) if indent_chars > 1 else 1)
        values = [_parse_number(c) for c in r[1:]]
        role = _classify_row(stripped, indent, stype)
        rows_out.append(StatementRow(label=stripped, values=values, indent=indent, role=role))

    # If the column labels are FULL dates (e.g., "December 31, 2026"), the
    # default period_header ("December 31,") is redundant — suppress it so the
    # renderer doesn't double-print the same as-of date language.
    labels_are_full_dates = any(
        re.match(r"[A-Za-z]+ ?\d{1,2}, ?\d{4}", lbl) for lbl in period_labels
    )
    period_header = "" if labels_are_full_dates else _DEFAULT_PERIOD_HEADER[stype]

    return StatementSegment(
        statement_type=stype,
        title=_TITLE_HINT_MAP[stype],
        units_note=_DEFAULT_UNITS[stype],
        period_header=period_header,
        period_labels=period_labels,
        rows=rows_out,
    )


def _from_csv(path: Path, delim: str) -> tuple[str, list[StatementSegment]]:
    raw_rows: list[list[Any]] = []
    text_parts: list[str] = []
    with path.open(newline="", encoding="utf-8") as fh:
        for r in csv.reader(fh, delimiter=delim):
            raw_rows.append(r)
            text_parts.extend(c for c in r if c)
    joined = "\n".join(text_parts)
    mode = _detect_mode(joined)
    segments: list[StatementSegment] = []
    if mode.startswith("segment:"):
        stype = mode.split(":", 1)[1]
        segments.append(_segment_from_rows(stype, raw_rows))
    return joined, segments


def _from_text(path: Path) -> tuple[str, list[StatementSegment]]:
    """Plain text / markdown — no segment extraction, full-filing only."""
    return path.read_text(encoding="utf-8", errors="replace"), []


def _from_docx(path: Path) -> tuple[str, list[StatementSegment]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit(f"python-docx required for docx input: {exc}")
    doc = Document(path)
    text_parts: list[str] = []
    raw_rows: list[list[Any]] = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            raw_rows.append(cells)
            text_parts.extend(cells)
    joined = "\n".join(text_parts)
    mode = _detect_mode(joined)
    segments: list[StatementSegment] = []
    if mode.startswith("segment:") and raw_rows:
        stype = mode.split(":", 1)[1]
        segments.append(_segment_from_rows(stype, raw_rows))
    return joined, segments


def _from_pdf(path: Path) -> tuple[str, list[StatementSegment]]:
    try:
        import pdfplumber
    except ImportError as exc:
        raise SystemExit(f"pdfplumber required for pdf input: {exc}")
    text_parts: list[str] = []
    raw_rows: list[list[Any]] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            text_parts.append(t)
            for tbl in page.extract_tables() or []:
                raw_rows.extend(tbl)
    joined = "\n".join(text_parts)
    mode = _detect_mode(joined)
    segments: list[StatementSegment] = []
    if mode.startswith("segment:") and raw_rows:
        stype = mode.split(":", 1)[1]
        segments.append(_segment_from_rows(stype, raw_rows))
    return joined, segments


def _from_json(path: Path) -> tuple[str, list[StatementSegment]]:
    data = json.loads(path.read_text(encoding="utf-8"))
    # If the JSON already has our schema, just round-trip.
    if isinstance(data, dict) and "mode_hint" in data and "segments" in data:
        segs = [
            StatementSegment(**{
                **{k: v for k, v in s.items() if k != "rows"},
                "rows": [StatementRow(**r) for r in s.get("rows", [])],
            })
            for s in data.get("segments", [])
        ]
        text = json.dumps(data, ensure_ascii=False)
        return text, segs
    text = json.dumps(data, ensure_ascii=False)
    return text, []


# ---------------------------------------------------------------------------
# Top-level dispatch
# ---------------------------------------------------------------------------

_EXT_DISPATCH = {
    ".xlsx": ("xlsx", _from_xlsx),
    ".xlsm": ("xlsx", _from_xlsx),
    ".csv":  ("csv",  lambda p: _from_csv(p, ",")),
    ".tsv":  ("tsv",  lambda p: _from_csv(p, "\t")),
    ".docx": ("docx", _from_docx),
    ".md":   ("md",   _from_text),
    ".txt":  ("txt",  _from_text),
    ".html": ("html", _from_text),
    ".htm":  ("html", _from_text),
    ".pdf":  ("pdf",  _from_pdf),
    ".json": ("json", _from_json),
}


def extract(path: Path) -> Extracted:
    ext = path.suffix.lower()
    if ext not in _EXT_DISPATCH:
        raise SystemExit(f"unsupported extension: {ext}")
    source_format, fn = _EXT_DISPATCH[ext]
    joined_text, segments = fn(path)
    if segments:
        # If we have segments, mode is segment:<first_type>; multi-segment is
        # treated as full_filing because it carries multiple statements.
        mode_hint = (
            f"segment:{segments[0].statement_type}"
            if len(segments) == 1
            else "full_filing"
        )
    else:
        mode_hint = _detect_mode(joined_text)
    return Extracted(
        mode_hint=mode_hint,
        source_format=source_format,
        metadata=FilingMetadata(),
        segments=segments,
        body_sections=[],
    )


def _to_jsonable(obj: Any) -> Any:
    if hasattr(obj, "__dataclass_fields__"):
        return {k: _to_jsonable(v) for k, v in asdict(obj).items()}
    if isinstance(obj, list):
        return [_to_jsonable(x) for x in obj]
    if isinstance(obj, dict):
        return {k: _to_jsonable(v) for k, v in obj.items()}
    return obj


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("input", type=Path, help="input file")
    ap.add_argument("--out", type=Path, default=None, help="output JSON path (default: stdout)")
    ap.add_argument("--mode", default=None, help="override mode_hint")
    ap.add_argument("--registrant", default=None)
    ap.add_argument("--filing-type", default=None)
    ap.add_argument("--period-end", default=None)
    ap.add_argument("--consolidated", action="store_true",
                    help="prefix titles with CONSOLIDATED (default off for Envision mocks)")
    args = ap.parse_args(argv)

    extracted = extract(args.input)
    if args.mode:
        extracted.mode_hint = args.mode
    if args.registrant:
        extracted.metadata.registrant = args.registrant
    if args.filing_type:
        extracted.metadata.filing_type = args.filing_type
    if args.period_end:
        extracted.metadata.period_end = args.period_end
    extracted.metadata.consolidated = args.consolidated
    if not args.consolidated:
        # Strip "CONSOLIDATED " prefix from segment titles
        for seg in extracted.segments:
            if seg.title.startswith("CONSOLIDATED "):
                seg.title = seg.title.replace("CONSOLIDATED ", "")

    payload = _to_jsonable(extracted)
    text = json.dumps(payload, indent=2, ensure_ascii=False)
    if args.out:
        args.out.write_text(text, encoding="utf-8")
    else:
        sys.stdout.write(text + "\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
