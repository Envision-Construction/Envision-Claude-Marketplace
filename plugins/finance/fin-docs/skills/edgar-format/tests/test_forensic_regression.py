#!/usr/bin/env python3
"""
test_forensic_regression.py — forensic accuracy regression test for edgar-format.

Generates synthetic IS / BS / CFS fixtures, runs the full extract → render
pipeline, then asserts every forensic invariant from
references/financial-statements.md. Fails LOUDLY on any deviation. Run before
shipping any change to extract.py or render_xlsx.py.

Each assertion has a `code` (used in failure messages) and an `issue` field so
when a test fails the user knows exactly which spec rule was violated. The
output ends with a one-line summary of PASS/FAIL counts.

Usage:
    python3 tests/test_forensic_regression.py

Exit code is non-zero if any assertion fails.
"""
from __future__ import annotations

import json
import re
import subprocess
import sys
from dataclasses import dataclass, field
from pathlib import Path
from tempfile import TemporaryDirectory

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Border, Font, Side


# ---------------------------------------------------------------------------
# Fixtures — synthetic statement xlsx files mimicking real uploads
# ---------------------------------------------------------------------------

def write_income_statement(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Income Statement"
    rows = [
        ["Envision Construction LLC Income Statement", None, None, None],
        ["For Years Ended December 31", None, None, None],
        [None, "2026", "2025", "2024"],
        ["Revenue", 1234567, 1100000, 950000],
        ["Cost of revenue", 800000, 720000, 640000],
        ["Gross profit", 434567, 380000, 310000],
        ["Operating expenses:", None, None, None],
        ["  Selling, general and administrative", 120000, 100000, 85000],
        ["  Research and development", 60000, 55000, 48000],
        ["  Depreciation and amortization", 22000, 20000, 18000],
        ["Total operating expenses", 202000, 175000, 151000],
        ["Income from operations", 232567, 205000, 159000],
        ["Interest expense", -12000, -10000, -8000],
        ["Other income, net", 2500, 1500, 1000],
        ["Income before income taxes", 223067, 196500, 152000],
        ["Income tax provision", -55767, -49125, -38000],
        ["Net income", 167300, 147375, 114000],
        [None, None, None, None],
        ["Earnings per common share:", None, None, None],
        ["  Basic", 1.45, 1.27, 0.98],
        ["  Diluted", 1.43, 1.25, 0.97],
        [None, None, None, None],
        ["Weighted-average common shares outstanding:", None, None, None],
        ["  Basic", 100544, 100346, 100510],
        ["  Diluted", 101943, 101952, 101547],
    ]
    for r in rows:
        ws.append(r)
    wb.save(path)


def write_balance_sheet(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Balance Sheet"
    rows = [
        ["Envision Construction LLC Balance Sheet", None, None],
        [None, "December 31, 2026", "December 31, 2025"],
        ["Assets", None, None],
        ["Current assets:", None, None],
        ["  Cash and cash equivalents", 245000, 180000],
        ["  Short-term investments", 50000, 45000],
        ["  Accounts receivable, net", 125000, 110000],
        ["  Inventory", 80000, 72000],
        ["Total current assets", 500000, 407000],
        ["Property and equipment, net", 340000, 310000],
        ["Goodwill", 120000, 120000],
        ["Total assets", 960000, 837000],
        ["Liabilities and stockholders' equity", None, None],
        ["Current liabilities:", None, None],
        ["  Accounts payable", 95000, 88000],
        ["  Accrued liabilities", 62000, 55000],
        ["Total current liabilities", 157000, 143000],
        ["Long-term debt, net", 180000, 200000],
        ["Total liabilities", 337000, 343000],
        ["Stockholders' equity:", None, None],
        ["  Common stock, $0.001 par value", 100, 99],
        ["  Additional paid-in capital", 315000, 310000],
        ["  Retained earnings", 307900, 183901],
        ["Total stockholders' equity", 623000, 494000],
        ["Total liabilities and stockholders' equity", 960000, 837000],
    ]
    for r in rows:
        ws.append(r)
    wb.save(path)


def write_cash_flows(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Cash Flows"
    rows = [
        ["Envision Construction LLC Statement of Cash Flows", None, None, None],
        [None, "2026", "2025", "2024"],
        ["Cash flows from operating activities:", None, None, None],
        ["  Net income", 167300, 147375, 114000],
        ["  Depreciation and amortization", 22000, 20000, 18000],
        ["  Stock-based compensation", 18000, 16000, 14000],
        ["Net cash provided by operating activities", 207300, 183375, 146000],
        ["Cash flows from investing activities:", None, None, None],
        ["  Purchases of property and equipment", -50000, -42000, -38000],
        ["  Purchases of investments", -60000, -50000, -40000],
        ["Net cash used in investing activities", -110000, -92000, -78000],
        ["Cash flows from financing activities:", None, None, None],
        ["  Repayments of long-term debt", -20000, -30000, -15000],
        ["  Repurchases of common stock", -40000, -30000, 0],
        ["Net cash provided by (used in) financing activities", -60000, -60000, -15000],
        ["Net increase in cash and cash equivalents", 37300, 31375, 53000],
        ["Cash and cash equivalents, beginning of period", 207700, 176325, 123325],
        ["Cash and cash equivalents, end of period", 245000, 207700, 176325],
        ["Supplemental disclosures of cash flow information:", None, None, None],
        ["  Cash paid for interest", 12500, 10200, 8100],
        ["  Cash paid for income taxes", 52000, 46000, 35000],
    ]
    for r in rows:
        ws.append(r)
    wb.save(path)


# ---------------------------------------------------------------------------
# Assertion framework — collect failures, report all of them
# ---------------------------------------------------------------------------

@dataclass
class TestRun:
    name: str
    passes: int = 0
    fails: list[str] = field(default_factory=list)

    def check(self, code: str, condition: bool, message: str) -> None:
        if condition:
            self.passes += 1
        else:
            self.fails.append(f"[{code}] {message}")

    def summary(self) -> str:
        if not self.fails:
            return f"  {self.name}: {self.passes} passed"
        out = [f"  {self.name}: {self.passes} passed, {len(self.fails)} FAILED"]
        out.extend(f"      ✗ {f}" for f in self.fails)
        return "\n".join(out)


SKILL_DIR = Path(__file__).resolve().parent.parent
SCRIPTS = SKILL_DIR / "scripts"


def run_pipeline(input_xlsx: Path, tmpdir: Path, registrant: str = "Envision Construction LLC") -> Path:
    """Run extract.py → render_xlsx.py and return path to rendered xlsx."""
    extracted = tmpdir / "extracted.json"
    rendered = tmpdir / "rendered.xlsx"
    subprocess.run(
        [
            sys.executable, str(SCRIPTS / "extract.py"), str(input_xlsx),
            "--registrant", registrant,
            "--filing-type", "10-K",
            "--out", str(extracted),
        ],
        check=True,
    )
    subprocess.run(
        [sys.executable, str(SCRIPTS / "render_xlsx.py"), str(extracted), str(rendered)],
        check=True,
    )
    return rendered


# ---------------------------------------------------------------------------
# Forensic invariants — derived from references/financial-statements.md
# ---------------------------------------------------------------------------

EXPECTED_FONT = "Times New Roman"
DOLLAR_FORMAT_FRAG = '$* '       # presence of "$* " marks a dollar-floating accounting format
PARENS_FORMAT_FRAG = '(#,##0'    # presence indicates negative-in-parens
EM_DASH_FORMAT_FRAG = '"—"'      # presence indicates em-dash for zero


def _has_top_thin(cell) -> bool:
    return bool(cell.border and cell.border.top and cell.border.top.style == "thin")


def _has_bottom_double(cell) -> bool:
    return bool(cell.border and cell.border.bottom and cell.border.bottom.style == "double")


def _find_row(ws, label_substring: str) -> int | None:
    for r in range(1, ws.max_row + 1):
        v = ws.cell(row=r, column=1).value
        if v and label_substring.lower() in str(v).lower():
            return r
    return None


def _all_data_cells(ws, row: int, first_col: int = 2):
    return [ws.cell(row=row, column=c) for c in range(first_col, ws.max_column + 1)
            if ws.cell(row=row, column=c).value is not None]


def assert_title_block(ws, registrant: str, expected_title: str, run: TestRun) -> None:
    a1 = ws.cell(row=1, column=1)
    run.check("TITLE-01", a1.value == registrant,
              f"row 1 col A expected registrant '{registrant}', got {a1.value!r}")
    run.check("TITLE-02", a1.font.name == EXPECTED_FONT,
              f"row 1 font should be {EXPECTED_FONT}, got {a1.font.name!r}")
    run.check("TITLE-03", a1.font.bold is True,
              "row 1 (registrant) should be bold")
    run.check("TITLE-04", a1.alignment.horizontal == "center",
              "row 1 (registrant) should be horizontally centered")

    a2 = ws.cell(row=2, column=1)
    run.check("TITLE-05", expected_title in str(a2.value or ""),
              f"row 2 should contain {expected_title!r}, got {a2.value!r}")
    run.check("TITLE-06", a2.font.bold is True,
              "row 2 (statement title) should be bold")
    run.check("TITLE-07", a2.value and a2.value == a2.value.upper(),
              "row 2 (statement title) should be ALL CAPS")

    a3 = ws.cell(row=3, column=1)
    run.check("TITLE-08", a3.value and "in " in str(a3.value).lower(),
              "row 3 should be units note (contains 'in thousands' or 'in millions')")
    run.check("TITLE-09", a3.font.italic is True,
              "row 3 (units note) should be italic")


def assert_column_headers(ws, run: TestRun) -> None:
    """Find the year/date header row, assert bold + right-aligned + bottom rule."""
    header_row = None
    for r in range(1, 12):
        b = ws.cell(row=r, column=2).value
        if b and (re.match(r"^(19|20)\d{2}$", str(b).strip())
                  or re.match(r"[A-Za-z]+ ?\d{1,2}, ?\d{4}", str(b))):
            header_row = r
            break
    run.check("HEAD-01", header_row is not None, "year/date header row not found")
    if header_row is None:
        return
    for col in range(2, ws.max_column + 1):
        c = ws.cell(row=header_row, column=col)
        if c.value is None:
            continue
        run.check("HEAD-02", c.font.bold is True,
                  f"col-header {c.coordinate} should be bold (value={c.value!r})")
        run.check("HEAD-03", c.alignment.horizontal == "right",
                  f"col-header {c.coordinate} should be right-aligned")
        run.check("HEAD-04",
                  c.border and c.border.bottom and c.border.bottom.style == "thin",
                  f"col-header {c.coordinate} should have a thin bottom rule")


def assert_universal_invariants(ws, run: TestRun) -> None:
    """Apply to every numeric data cell."""
    for row in ws.iter_rows(min_row=4, max_row=ws.max_row):
        for c in row:
            if isinstance(c.value, (int, float)) and not isinstance(c.value, bool):
                run.check("UNI-01", c.font.name == EXPECTED_FONT,
                          f"{c.coordinate} font should be Times New Roman, got {c.font.name!r}")
                run.check("UNI-02", abs((c.font.size or 0) - 10) < 0.01,
                          f"{c.coordinate} font size should be 10, got {c.font.size}")
                run.check("UNI-03", c.alignment.horizontal == "right",
                          f"{c.coordinate} numeric should be right-aligned")
                run.check("UNI-04", PARENS_FORMAT_FRAG in c.number_format,
                          f"{c.coordinate} format should render negatives in parens; got {c.number_format!r}")
                run.check("UNI-05", EM_DASH_FORMAT_FRAG in c.number_format,
                          f"{c.coordinate} format should render zero as em-dash; got {c.number_format!r}")
                run.check("UNI-06", c.font.bold is False,
                          f"{c.coordinate} body number should NOT be bold (EDGAR uses rules for emphasis, not weight)")


def assert_floating_dollar_pattern(ws, run: TestRun, *, expect_first_data_label: str,
                                   subtotal_labels: list[str], grand_total_labels: list[str],
                                   mid_data_labels: list[str]) -> None:
    """The floating $ shows on (a) first data row and (b) every subtotal/grand-total row.
    Middle data rows must NOT carry a $-floating format."""
    first = _find_row(ws, expect_first_data_label)
    run.check("DLR-01", first is not None,
              f"could not find first-data-row labeled {expect_first_data_label!r}")
    if first is not None:
        for c in _all_data_cells(ws, first):
            run.check("DLR-02", DOLLAR_FORMAT_FRAG in c.number_format,
                      f"first data row {c.coordinate} ({expect_first_data_label}) should float $; got {c.number_format!r}")

    for label in subtotal_labels:
        r = _find_row(ws, label)
        run.check("DLR-03", r is not None, f"could not find subtotal row labeled {label!r}")
        if r is None:
            continue
        for c in _all_data_cells(ws, r):
            run.check("DLR-04", DOLLAR_FORMAT_FRAG in c.number_format,
                      f"subtotal {c.coordinate} ({label}) should float $; got {c.number_format!r}")
            run.check("DLR-05", _has_top_thin(c),
                      f"subtotal {c.coordinate} ({label}) needs thin top border")

    for label in grand_total_labels:
        r = _find_row(ws, label)
        run.check("DLR-06", r is not None, f"could not find grand-total row labeled {label!r}")
        if r is None:
            continue
        for c in _all_data_cells(ws, r):
            run.check("DLR-07", DOLLAR_FORMAT_FRAG in c.number_format,
                      f"grand_total {c.coordinate} ({label}) should float $; got {c.number_format!r}")
            run.check("DLR-08", _has_top_thin(c),
                      f"grand_total {c.coordinate} ({label}) needs thin top border")
            run.check("DLR-09", _has_bottom_double(c),
                      f"grand_total {c.coordinate} ({label}) needs double bottom border")

    for label in mid_data_labels:
        r = _find_row(ws, label)
        if r is None:
            continue  # tolerated — label might not exist in this fixture
        for c in _all_data_cells(ws, r):
            run.check("DLR-10", DOLLAR_FORMAT_FRAG not in c.number_format,
                      f"middle data row {c.coordinate} ({label}) must NOT float $; got {c.number_format!r}")


def assert_eps_vs_shares_disambiguation(ws, run: TestRun) -> None:
    """Under 'Earnings per common share:' Basic/Diluted use 2-decimal format.
    Under 'Weighted-average ... shares outstanding:' Basic/Diluted use integer format."""
    eps_section = _find_row(ws, "Earnings per common share")
    shares_section = _find_row(ws, "Weighted-average common shares")
    if eps_section is None or shares_section is None:
        return  # not applicable to this fixture
    # Rows immediately after eps_section (skipping spacers) are Basic/Diluted EPS
    for r in range(eps_section + 1, shares_section):
        if ws.cell(row=r, column=2).value is None:
            continue
        for c in _all_data_cells(ws, r):
            run.check("EPS-01", "0.00" in c.number_format,
                      f"EPS row {c.coordinate} should use 2-decimal format; got {c.number_format!r}")
    for r in range(shares_section + 1, ws.max_row + 1):
        if ws.cell(row=r, column=2).value is None:
            continue
        for c in _all_data_cells(ws, r):
            run.check("EPS-02", "0.00" not in c.number_format,
                      f"shares row {c.coordinate} should be integer format; got {c.number_format!r}")


def assert_page_setup(ws, run: TestRun) -> None:
    run.check("PAGE-01", abs(ws.page_margins.top - 1.0) < 0.01, "top margin should be 1.0 inch")
    run.check("PAGE-02", abs(ws.page_margins.bottom - 1.0) < 0.01, "bottom margin should be 1.0 inch")
    run.check("PAGE-03", abs(ws.page_margins.left - 0.75) < 0.01, "left margin should be 0.75 inch")
    run.check("PAGE-04", abs(ws.page_margins.right - 0.75) < 0.01, "right margin should be 0.75 inch")
    run.check("PAGE-05", str(ws.page_setup.paperSize) == "1", "paper size should be Letter (1)")


# ---------------------------------------------------------------------------
# Per-statement test orchestration
# ---------------------------------------------------------------------------

def test_income_statement(tmpdir: Path) -> TestRun:
    run = TestRun("Income Statement")
    src = tmpdir / "in_is.xlsx"
    write_income_statement(src)
    out = run_pipeline(src, tmpdir)
    wb = load_workbook(out)
    run.check("SHEET-01", "IS" in wb.sheetnames,
              f"IS sheet missing; sheets = {wb.sheetnames}")
    ws = wb["IS"]

    assert_title_block(ws, "Envision Construction LLC", "STATEMENTS OF OPERATIONS", run)
    assert_column_headers(ws, run)
    assert_universal_invariants(ws, run)
    assert_floating_dollar_pattern(
        ws, run,
        expect_first_data_label="Revenue",
        subtotal_labels=[
            "Gross profit", "Total operating expenses",
            "Income from operations", "Income before income taxes",
        ],
        grand_total_labels=["Net income"],
        mid_data_labels=[
            "Cost of revenue", "Selling, general and administrative",
            "Interest expense", "Income tax provision",
        ],
    )
    assert_eps_vs_shares_disambiguation(ws, run)
    assert_page_setup(ws, run)
    return run


def test_balance_sheet(tmpdir: Path) -> TestRun:
    run = TestRun("Balance Sheet")
    src = tmpdir / "in_bs.xlsx"
    write_balance_sheet(src)
    out = run_pipeline(src, tmpdir)
    wb = load_workbook(out)
    run.check("SHEET-02", "BS" in wb.sheetnames, f"BS sheet missing; sheets = {wb.sheetnames}")
    ws = wb["BS"]

    assert_title_block(ws, "Envision Construction LLC", "BALANCE SHEETS", run)
    assert_column_headers(ws, run)
    assert_universal_invariants(ws, run)
    assert_floating_dollar_pattern(
        ws, run,
        expect_first_data_label="Cash and cash equivalents",
        subtotal_labels=[
            "Total current assets", "Total current liabilities",
            "Total liabilities", "Total stockholders' equity",
        ],
        # Balance sheet has TWO grand totals — Total assets, Total liabilities and stockholders' equity
        grand_total_labels=["Total assets", "Total liabilities and stockholders' equity"],
        mid_data_labels=["Inventory", "Goodwill", "Accrued liabilities"],
    )

    # BS-specific: section banners must be ALL CAPS uppercased
    assets_row = _find_row(ws, "ASSETS")
    run.check("BS-01", assets_row is not None, "ASSETS banner row not found (expected uppercased)")
    if assets_row:
        run.check("BS-02", ws.cell(row=assets_row, column=1).value == "ASSETS",
                  f"ASSETS banner should be exact uppercase; got {ws.cell(row=assets_row, column=1).value!r}")

    # Balance check — the two grand totals must be numerically equal
    ta = _find_row(ws, "Total assets")
    tle = _find_row(ws, "Total liabilities and stockholders' equity")
    if ta is not None and tle is not None:
        for col in range(2, ws.max_column + 1):
            v1 = ws.cell(row=ta, column=col).value
            v2 = ws.cell(row=tle, column=col).value
            if v1 is None and v2 is None:
                continue
            run.check("BS-03", v1 == v2,
                      f"balance check: Total assets ({v1}) must equal Total L&E ({v2}) in col {col}")

    assert_page_setup(ws, run)
    return run


def test_cash_flows(tmpdir: Path) -> TestRun:
    run = TestRun("Cash Flows")
    src = tmpdir / "in_cf.xlsx"
    write_cash_flows(src)
    out = run_pipeline(src, tmpdir)
    wb = load_workbook(out)
    run.check("SHEET-03", "CF" in wb.sheetnames, f"CF sheet missing; sheets = {wb.sheetnames}")
    ws = wb["CF"]

    assert_title_block(ws, "Envision Construction LLC", "STATEMENTS OF CASH FLOWS", run)
    assert_column_headers(ws, run)
    assert_universal_invariants(ws, run)
    assert_floating_dollar_pattern(
        ws, run,
        expect_first_data_label="Net income",
        subtotal_labels=[
            "Net cash provided by operating activities",
            "Net cash used in investing activities",
            "Net cash provided by (used in) financing activities",
            "Net increase in cash and cash equivalents",
        ],
        grand_total_labels=["Cash and cash equivalents, end of period"],
        mid_data_labels=[
            "Depreciation and amortization", "Stock-based compensation",
            "Purchases of property and equipment", "Repayments of long-term debt",
        ],
    )

    # CFS-specific: three section headers exist
    for header in (
        "Cash flows from operating activities:",
        "Cash flows from investing activities:",
        "Cash flows from financing activities:",
    ):
        run.check("CFS-01", _find_row(ws, header) is not None,
                  f"missing CFS section header {header!r}")

    assert_page_setup(ws, run)
    return run


# ---------------------------------------------------------------------------
# Detection regression — ambiguity, multi-sheet, non-financial fallthrough
# ---------------------------------------------------------------------------

def test_detection_invariants(tmpdir: Path) -> TestRun:
    run = TestRun("Detection")

    # Income statement → segment:income_statement
    src = tmpdir / "det_is.xlsx"
    write_income_statement(src)
    extracted = tmpdir / "det_is.json"
    subprocess.run(
        [sys.executable, str(SCRIPTS / "extract.py"), str(src),
         "--registrant", "X", "--out", str(extracted)],
        check=True,
    )
    data = json.loads(extracted.read_text())
    run.check("DET-01", data["mode_hint"] == "segment:income_statement",
              f"IS fixture should detect as segment:income_statement, got {data['mode_hint']}")

    # Balance sheet
    src = tmpdir / "det_bs.xlsx"
    write_balance_sheet(src)
    extracted = tmpdir / "det_bs.json"
    subprocess.run(
        [sys.executable, str(SCRIPTS / "extract.py"), str(src),
         "--registrant", "X", "--out", str(extracted)],
        check=True,
    )
    data = json.loads(extracted.read_text())
    run.check("DET-02", data["mode_hint"] == "segment:balance_sheet",
              f"BS fixture should detect as segment:balance_sheet, got {data['mode_hint']}")

    # Cash flows
    src = tmpdir / "det_cf.xlsx"
    write_cash_flows(src)
    extracted = tmpdir / "det_cf.json"
    subprocess.run(
        [sys.executable, str(SCRIPTS / "extract.py"), str(src),
         "--registrant", "X", "--out", str(extracted)],
        check=True,
    )
    data = json.loads(extracted.read_text())
    run.check("DET-03", data["mode_hint"] == "segment:cash_flows",
              f"CF fixture should detect as segment:cash_flows, got {data['mode_hint']}")

    # Non-financial document → full_filing (no signature hits)
    src = tmpdir / "narrative.md"
    src.write_text("# About Envision\n\nWe build things.\nOur mission is to ship.\n")
    extracted = tmpdir / "narrative.json"
    subprocess.run(
        [sys.executable, str(SCRIPTS / "extract.py"), str(src), "--out", str(extracted)],
        check=True,
    )
    data = json.loads(extracted.read_text())
    run.check("DET-04", data["mode_hint"] == "full_filing",
              f"narrative.md should detect as full_filing, got {data['mode_hint']}")

    return run


# ---------------------------------------------------------------------------
# HTML / DOCX renderer regression — same forensic spec, different format
# ---------------------------------------------------------------------------

def test_html_rendering(tmpdir: Path) -> TestRun:
    """Verify render_html.py emits the expected EDGAR semantic classes and
    content so the CSS contract in assets/edgar.css can be relied upon."""
    run = TestRun("HTML")
    src = tmpdir / "html_is.xlsx"
    write_income_statement(src)
    extracted = tmpdir / "html_is.json"
    subprocess.run(
        [sys.executable, str(SCRIPTS / "extract.py"), str(src),
         "--registrant", "Envision Construction LLC",
         "--filing-type", "10-K",
         "--out", str(extracted)],
        check=True,
    )
    out_html = tmpdir / "is.html"
    subprocess.run(
        [sys.executable, str(SCRIPTS / "render_html.py"), str(extracted), str(out_html)],
        check=True,
    )
    body = out_html.read_text(encoding="utf-8")

    # Structural assertions — same forensic invariants, expressed in HTML/CSS terms
    run.check("HTML-01", "<!doctype html>" in body.lower(), "HTML output should declare doctype")
    run.check("HTML-02", "Times New Roman" in body, "Times New Roman should appear in inline CSS")
    run.check("HTML-03", 'class="edgar-registrant"' in body, "registrant element missing")
    run.check("HTML-04", 'class="edgar-title"' in body, "title element missing")
    run.check("HTML-05", 'class="edgar-units"' in body, "units note missing")
    run.check("HTML-06", "STATEMENTS OF OPERATIONS" in body, "statement title should be uppercased")
    run.check("HTML-07", 'class="edgar-col-year"' in body, "column-year header class missing")
    run.check("HTML-08", "edgar-subtotal" in body, "subtotal class missing")
    run.check("HTML-09", "edgar-grand-total" in body, "grand-total class missing")
    run.check("HTML-10", "edgar-dollar" in body, "floating-dollar class missing")
    run.check("HTML-11", "edgar-section-header" in body, "section_header class missing")
    run.check("HTML-12", "(In thousands, except per share data)" in body, "units note text missing")
    # Negative number rendered as parens, not -
    run.check("HTML-13", "(12,000)" in body, "negative should render as (12,000) not -12000")
    # Em-dash not rendered for IS values (they're all numeric)
    run.check("HTML-14", "Revenue" in body and "1,234,567" in body, "Revenue 1,234,567 should appear")
    # EPS rendered with 2 decimals
    run.check("HTML-15", "1.45" in body and "1.43" in body, "EPS values should render with 2 decimals")
    # Share counts rendered as integers (no decimals)
    run.check("HTML-16", "100,544" in body, "share counts should render as integers with thousands sep")
    # CSS contract: no inline color overrides on numeric cells
    run.check("HTML-17", "color: red" not in body and "color:red" not in body,
              "no color overrides on numeric cells")

    return run


def test_docx_rendering(tmpdir: Path) -> TestRun:
    """Verify render_docx.py applies Times New Roman + EDGAR conventions to
    a docx output. Reads the docx back via python-docx and asserts."""
    run = TestRun("DOCX")
    src = tmpdir / "docx_is.xlsx"
    write_income_statement(src)
    extracted = tmpdir / "docx_is.json"
    subprocess.run(
        [sys.executable, str(SCRIPTS / "extract.py"), str(src),
         "--registrant", "Envision Construction LLC",
         "--out", str(extracted)],
        check=True,
    )
    out_docx = tmpdir / "is.docx"
    subprocess.run(
        [sys.executable, str(SCRIPTS / "render_docx.py"), str(extracted), str(out_docx)],
        check=True,
    )

    try:
        from docx import Document
    except ImportError:
        run.check("DOCX-00", False, "python-docx not installed — skipping docx assertions")
        return run

    doc = Document(out_docx)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    # Concatenate table text too
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text

    run.check("DOCX-01", doc.styles["Normal"].font.name == "Times New Roman",
              f"Normal style font should be Times New Roman, got {doc.styles['Normal'].font.name!r}")
    run.check("DOCX-02", doc.styles["Normal"].font.size is not None and doc.styles["Normal"].font.size.pt == 10,
              "Normal style font size should be 10pt")
    run.check("DOCX-03", "Envision Construction LLC" in all_text, "registrant name missing")
    run.check("DOCX-04", "STATEMENTS OF OPERATIONS" in all_text, "ALL-CAPS statement title missing")
    run.check("DOCX-05", "(In thousands, except per share data)" in all_text, "units note missing")
    run.check("DOCX-06", "Revenue" in all_text, "Revenue label missing")
    run.check("DOCX-07", "1,234,567" in all_text, "Revenue value with thousands sep missing")
    run.check("DOCX-08", "(12,000)" in all_text, "negative in parens missing")
    run.check("DOCX-09", "—" in all_text or "1.45" in all_text, "EPS or em-dash missing")

    # Page margins per spec
    sec = doc.sections[0]
    run.check("DOCX-10", abs(sec.top_margin.inches - 1.0) < 0.01, "top margin should be 1.0in")
    run.check("DOCX-11", abs(sec.left_margin.inches - 0.75) < 0.01, "left margin should be 0.75in")

    # Confirm at least one cell has a top border (the subtotal/grand_total rule)
    has_top_border = False
    has_double_bottom = False
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                tcPr = cell._tc.tcPr
                if tcPr is None:
                    continue
                borders = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
                if borders is None:
                    continue
                top = borders.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}top')
                bot = borders.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom')
                if top is not None and top.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') in {"single"}:
                    has_top_border = True
                if bot is not None and bot.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') == "double":
                    has_double_bottom = True
    run.check("DOCX-12", has_top_border, "no cells with a single top border — subtotals/grand_totals missing rule")
    run.check("DOCX-13", has_double_bottom, "no cells with a double bottom border — grand_total rule missing")

    return run


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    print("=== edgar-format forensic regression suite ===\n")
    runs: list[TestRun] = []
    with TemporaryDirectory() as td:
        tmpdir = Path(td)
        runs.append(test_income_statement(tmpdir))
        runs.append(test_balance_sheet(tmpdir))
        runs.append(test_cash_flows(tmpdir))
        runs.append(test_detection_invariants(tmpdir))
        runs.append(test_html_rendering(tmpdir))
        runs.append(test_docx_rendering(tmpdir))

    total_pass = sum(r.passes for r in runs)
    total_fail = sum(len(r.fails) for r in runs)

    for r in runs:
        print(r.summary())

    print()
    print(f"=== TOTAL: {total_pass} passed, {total_fail} failed ===")
    if total_fail:
        print("\nForensic accuracy regression — fix renderer or extractor before shipping.")
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
